from __future__ import annotations

import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from work_agent_core.docx_exporter import create_docx, resolve_document_title
from work_agent_core.progress import command_heartbeat_text
from work_agent_core.office_preview import find_soffice
from work_agent_core.shell_tools import ShellExecutionTools, safe_environment
from work_agent_core.skill_manifest import office_python, probe_skill_environment
from work_agent_core.skill_runtime import render_skill_tool_arguments
from work_agent_core import web_server


class RuntimeEnvironmentTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.workspace_root = Path(__file__).resolve().parents[1]
        cls.agent_python = cls.workspace_root / ".venv" / "bin" / "python"
        if not cls.agent_python.is_file():
            raise unittest.SkipTest("project agent environment is not initialized")

    def test_safe_shell_environment_prefers_project_venv(self) -> None:
        environment = safe_environment(self.workspace_root)

        self.assertEqual(environment["VIRTUAL_ENV"], str(self.workspace_root / ".venv"))
        self.assertEqual(environment["WORK_AGENT_PYTHON"], str(self.agent_python))
        self.assertEqual(environment["PATH"].split(":", 1)[0], str(self.agent_python.parent))
        self.assertEqual(environment["PIP_REQUIRE_VIRTUALENV"], "true")

    def test_bare_python_command_uses_project_venv(self) -> None:
        result = ShellExecutionTools(self.workspace_root).execute(
            {
                "command": "python3 -c \"import sys; print(sys.executable)\"",
                "approved_by_user": True,
            }
        )
        payload = json.loads(result)

        self.assertTrue(payload["ok"])
        self.assertIn(str(self.workspace_root / ".venv" / "bin"), payload["stdout"].strip())

    def test_managed_environment_exposes_existing_media_and_node_tools(self) -> None:
        tools = ShellExecutionTools(self.workspace_root)
        for binary in ("ffmpeg", "node"):
            payload = json.loads(tools.execute({"command": f"which {binary}"}))
            self.assertTrue(payload["ok"], payload)
            self.assertEqual(payload["permission"], "allow")

    def test_skill_precheck_verifies_declared_runtime_dependencies(self) -> None:
        for skill_id in ("meeting-minutes", "anysearch"):
            payload = probe_skill_environment(self.workspace_root, skill_id=skill_id)
            probe = payload["skills"][0]
            self.assertTrue(probe["declared"])
            self.assertTrue(probe["ready"], probe)
        meeting = probe_skill_environment(
            self.workspace_root, skill_id="meeting-minutes"
        )["skills"][0]
        self.assertTrue(meeting["paths"][".venv/bin/python"])

    def test_meeting_minutes_skill_has_working_meeting_reference_and_type_routing(self) -> None:
        skill_dir = self.workspace_root / "meeting_audio_minutes" / "skills" / "meeting-minutes"
        skill_text = (skill_dir / "SKILL.md").read_text(encoding="utf-8")
        reference_text = (skill_dir / "references" / "work-submission-writing.md").read_text(
            encoding="utf-8"
        )

        self.assertIn("项目推进、方案讨论、工作协调、一般来访座谈类", skill_text)
        self.assertIn("conclusion-led", skill_text)
        self.assertIn("不能作为切换到四段式研判模板的依据", reference_text)
        self.assertIn("保留可靠来源中的面积、数量、资源清单", reference_text)
        self.assertTrue((self.workspace_root / "meeting_audio_minutes" / "meeting_minutes_spec.md").is_file())
        self.assertIn("official-document", skill_text)
        self.assertIn("complete Word workflow", skill_text)
        self.assertNotIn("create_docx_from_markdown", skill_text)

    def test_official_document_routing_and_blank_open_source_format_default(self) -> None:
        self.assertTrue(web_server.looks_like_official_document_request("帮我起草一份请示"))
        self.assertTrue(web_server.looks_like_official_document_request("按公文格式排版"))
        self.assertFalse(web_server.looks_like_official_document_request("制作一份产品宣传册"))

        default_format = web_server.DEFAULT_AGENT_SETTINGS["company_document_format"]
        self.assertEqual(default_format, "")

    def test_skill_without_dependency_declaration_is_not_reported_ready(self) -> None:
        payload = probe_skill_environment(self.workspace_root, skill_id="skill-creator")
        probe = payload["skills"][0]
        self.assertFalse(probe["declared"])
        self.assertFalse(probe["ready"])

    def test_docx_runtime_has_validation_and_preview_dependencies(self) -> None:
        probe = probe_skill_environment(self.workspace_root, skill_id="docx")["skills"][0]
        self.assertTrue(probe["python_modules"]["defusedxml"])
        self.assertTrue(probe["binaries"]["soffice"])
        self.assertTrue(probe["ready"], probe)
        self.assertTrue(find_soffice().is_file())

    def test_boolean_skill_arguments_render_as_switches(self) -> None:
        rendered = render_skill_tool_arguments(
            [
                {"param": "auto_repair", "flag": "--auto-repair", "optional": True},
                {"param": "verbose", "flag": "-v", "optional": True},
                "{{path}}",
            ],
            {"auto_repair": True, "verbose": False, "path": "example.docx"},
        )
        self.assertEqual(rendered, ["--auto-repair", "example.docx"])

    def test_generated_docx_has_schema_complete_zoom_setting(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "minutes.docx"
            create_docx("# 测试会议纪要\n\n正文。\n", output)
            with zipfile.ZipFile(output) as archive:
                settings = archive.read("word/settings.xml").decode("utf-8")
        self.assertIn('w:percent="100"', settings)

    def test_generated_docx_supports_compact_inline_conclusion_paragraphs(self) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        markdown = (
            "# 产教融合公共实训基地建设交流会议纪要\n\n"
            "围绕机器人训练场设备和空间资源，研究建设产教融合公共实训基地。\n\n"
            "**一、形成资源互补。**示例单位提供机器人本体和场地；合作方提供课程与运营资源。\n"
        )
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "minutes.docx"
            create_docx(markdown, output)
            document = Document(output)

        section = document.sections[0]
        self.assertAlmostEqual(section.top_margin.cm, 1.5, places=1)
        self.assertAlmostEqual(section.bottom_margin.cm, 1.5, places=1)
        self.assertAlmostEqual(section.left_margin.cm, 2.2, places=1)
        self.assertAlmostEqual(section.right_margin.cm, 2.0, places=1)
        conclusion = document.paragraphs[2]
        self.assertEqual(conclusion.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertEqual(conclusion.paragraph_format.first_line_indent.twips, 640)
        self.assertEqual(conclusion.runs[0].text, "一、形成资源互补。")
        self.assertEqual(conclusion.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")), "黑体")
        self.assertTrue(conclusion.runs[1].text.startswith("示例单位提供"))

    def test_filename_style_date_prefix_does_not_duplicate_markdown_title(self) -> None:
        markdown = "# 示例科技来访座谈沟通会议纪要\n\n正文。\n"

        self.assertEqual(
            resolve_document_title(markdown, "0721示例科技来访座谈沟通会议纪要"),
            "示例科技来访座谈沟通会议纪要",
        )

        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "minutes.docx"
            create_docx(
                markdown,
                output,
                title="0721示例科技来访座谈沟通会议纪要",
            )
            document = Document(output)

        self.assertEqual(document.paragraphs[0].text, "示例科技来访座谈沟通会议纪要")
        self.assertEqual(
            sum(paragraph.text == "示例科技来访座谈沟通会议纪要" for paragraph in document.paragraphs),
            1,
        )

    def test_office_skills_use_project_venv(self) -> None:
        self.assertEqual(office_python(), self.agent_python)

    def test_command_heartbeat_leaves_animation_to_the_frontend(self) -> None:
        heartbeat = command_heartbeat_text("本地 Qwen3-ASR 转写", 20)

        self.assertEqual(heartbeat, "[20s] 本地 Qwen3-ASR 转写处理中...\n")
        self.assertNotIn("/", heartbeat)
        self.assertNotIn("\\", heartbeat)


if __name__ == "__main__":
    unittest.main()
