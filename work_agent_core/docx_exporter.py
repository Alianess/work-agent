from __future__ import annotations

from pathlib import Path
import argparse
import json
import re


BODY_EAST_ASIA_FONT = "仿宋_GB2312"
HEADING_EAST_ASIA_FONT = "黑体"
TITLE_EAST_ASIA_FONT = "方正小标宋简体"


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    if args.markdown_path:
        markdown = Path(args.markdown_path).read_text(encoding="utf-8", errors="replace")
    else:
        markdown = args.markdown_content or ""
    if not markdown.strip():
        raise ValueError("Markdown 内容为空。")
    output_path = Path(args.output_path).expanduser().resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    create_docx(markdown, output_path, title=args.title or "")
    print(
        json.dumps(
            {
                "output_path": str(output_path),
                "title": args.title or infer_title(markdown),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


def parse_args(argv: list[str] | None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Create a formal Chinese meeting-minutes DOCX from Markdown.")
    parser.add_argument("--markdown-path")
    parser.add_argument("--markdown-content")
    parser.add_argument("--output-path", required=True)
    parser.add_argument("--title")
    return parser.parse_args(argv)


def create_docx(markdown: str, output_path: Path, *, title: str = "") -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        # The bundled OOXML validator requires percent even when the template
        # selects bestFit. Word accepts the explicit value and validation stays
        # deterministic across python-docx template versions.
        zoom.set(qn("w:percent"), "100")
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_EAST_ASIA_FONT)
    normal.font.size = Pt(16)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(29)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(32)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_EAST_ASIA_FONT)
        style.font.size = Pt(16)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(29)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.first_line_indent = Pt(32)

    lines = markdown.splitlines()
    doc_title = resolve_document_title(markdown, title)
    if doc_title:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(31)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(strip_heading_marker(doc_title))
        run.bold = False
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), TITLE_EAST_ASIA_FONT)
        run.font.size = Pt(22)

    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()
        if not line:
            index += 1
            continue
        if line == "---":
            add_separator(document)
            index += 1
            continue
        if is_table_start(lines, index):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(document, table_lines, WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, qn)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            if text != strip_heading_marker(doc_title):
                add_section_heading(document, text, qn)
            index += 1
            continue
        if is_chinese_section_heading(line):
            add_section_heading(document, line, qn)
            index += 1
            continue
        unordered = re.match(r"^[-*]\s+(.+)$", line)
        if unordered:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, unordered.group(1))
            index += 1
            continue
        numbered = re.match(r"^\d+[.、]\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, numbered.group(1))
            index += 1
            continue
        inline_section = parse_inline_section_paragraph(line)
        if inline_section is not None:
            heading_text, body_text = inline_section
            add_inline_section_paragraph(document, heading_text, body_text, qn)
            index += 1
            continue
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Pt(32)
        add_inline_runs(paragraph, line)
        index += 1

    document.save(output_path)


def add_section_heading(document, text: str, qn) -> None:
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(32)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(29)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = False
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_EAST_ASIA_FONT)
    run.font.size = Pt(16)


def parse_inline_section_paragraph(text: str) -> tuple[str, str] | None:
    match = re.match(
        r"^(?:\*\*)?([一二三四五六七八九十]+、[^。；：]{2,24}。)(?:\*\*)?(.+)$",
        text.strip(),
    )
    if match is None or not match.group(2).strip():
        return None
    return match.group(1).strip(), match.group(2).strip()


def add_inline_section_paragraph(document, heading_text: str, body_text: str, qn) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(32)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(29)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(0)
    heading_run = paragraph.add_run(heading_text)
    heading_run.bold = False
    heading_run.font.name = "Times New Roman"
    heading_run._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_EAST_ASIA_FONT)
    heading_run.font.size = Pt(16)
    add_inline_runs(paragraph, body_text)


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_separator(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = 6
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9D9D9")
    border.append(bottom)
    p_pr.append(border)


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and lines[index].strip().startswith("|")
        and re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", lines[index + 1].strip()) is not None
    )


def add_markdown_table(document, table_lines: list[str], table_alignment, cell_vertical_alignment, qn) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    rows = [parse_table_row(line) for line in table_lines]
    if len(rows) >= 2 and all(re.match(r"^:?-{3,}:?$", cell.strip()) for cell in rows[1]):
        rows.pop(1)
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    table.alignment = table_alignment.LEFT
    for row_index, row_values in enumerate(rows):
        cells = table.rows[row_index].cells
        for column_index, value in enumerate(row_values):
            cell = cells[column_index]
            cell.vertical_alignment = cell_vertical_alignment.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value.strip())
            run.font.size = Pt(10.5)
            if row_index == 0:
                run.bold = True
                set_cell_shading(cell, "F2F2F2")
    document.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def infer_title(markdown: str) -> str:
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped:
            return strip_heading_marker(stripped)
    return "会议纪要"


def resolve_document_title(markdown: str, requested_title: str) -> str:
    markdown_title = infer_title(markdown)
    requested = strip_heading_marker(requested_title.strip())
    if not requested:
        return markdown_title
    # Callers sometimes pass the filename-style title, e.g. `0721会议名称`,
    # while the Markdown H1 correctly contains only the display title. Avoid
    # rendering both as separate headings in that common meeting-minutes case.
    without_date_prefix = re.sub(r"^\d{4}[\s._-]*", "", requested).strip()
    if without_date_prefix == markdown_title:
        return markdown_title
    return requested


def strip_heading_marker(text: str) -> str:
    return re.sub(r"^#{1,6}\s+", "", text).strip()


def is_chinese_section_heading(text: str) -> bool:
    if len(text) > 28:
        return False
    return re.match(r"^[一二三四五六七八九十]+、[^，。；：:,.]{2,24}$", text.strip()) is not None


if __name__ == "__main__":
    raise SystemExit(main())
